"""Pull plain text out of an uploaded resume.

Handles PDF and DOCX. Scanned/image-only PDFs are detected and reported rather
than silently producing an empty extraction — OCR is a later concern, but
telling the user plainly is not.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import List

log = logging.getLogger(__name__)

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
# A resume is a couple of pages. Anything longer is a thesis or an attack, and
# either way it should not reach the model.
MAX_CHARS = 40_000
MAX_PDF_PAGES = 10
# Under this much text, a PDF is almost certainly a scan of a page image.
MIN_TEXT_CHARS = 120

PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"  # DOCX is a zip container


class UnsupportedFile(Exception):
    """The upload isn't a resume we can read. The message is user-facing."""


@dataclass
class ExtractedText:
    text: str
    page_count: int = 0
    warnings: List[str] = field(default_factory=list)


def sniff_kind(data: bytes, filename: str) -> str:
    """Identify the file by content, falling back to the extension.

    Content first: the extension is attacker-controlled, the magic bytes are
    what the parser will actually act on.
    """
    if data.startswith(PDF_MAGIC):
        return "pdf"
    if data.startswith(ZIP_MAGIC):
        return "docx"
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return "pdf"
    if lowered.endswith(".docx"):
        return "docx"
    if lowered.endswith(".doc"):
        raise UnsupportedFile(
            "Old .doc files aren't supported — please save as PDF or .docx."
        )
    raise UnsupportedFile("Please upload a PDF or Word (.docx) file.")


def extract_text(data: bytes, filename: str) -> ExtractedText:
    if not data:
        raise UnsupportedFile("That file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise UnsupportedFile("That file is too large — the limit is 5 MB.")

    kind = sniff_kind(data, filename)
    result = _extract_pdf(data) if kind == "pdf" else _extract_docx(data)

    result.text = _tidy(result.text)
    if len(result.text) > MAX_CHARS:
        result.text = result.text[:MAX_CHARS]
        result.warnings.append("Only the first part of this document was read.")

    if len(result.text) < MIN_TEXT_CHARS:
        raise UnsupportedFile(
            "We couldn't read any text from that file. If it's a scan or a "
            "photo, try exporting a text-based PDF — or fill the form in "
            "yourself, it only takes a few minutes."
        )
    return result


def _extract_pdf(data: bytes) -> ExtractedText:
    try:
        import pymupdf
    except ImportError:  # pragma: no cover - dependency is declared
        raise UnsupportedFile("PDF support isn't available on this server.") from None

    try:
        with pymupdf.open(stream=data, filetype="pdf") as doc:
            if doc.needs_pass:
                raise UnsupportedFile(
                    "That PDF is password-protected. Please upload an "
                    "unlocked copy."
                )
            pages = min(doc.page_count, MAX_PDF_PAGES)
            chunks = [doc[i].get_text("text") for i in range(pages)]
            warnings = []
            if doc.page_count > pages:
                warnings.append(f"Only the first {pages} pages were read.")
            return ExtractedText(text="\n".join(chunks),
                                 page_count=doc.page_count, warnings=warnings)
    except UnsupportedFile:
        raise
    except Exception as exc:  # noqa: BLE001 - any parser failure is user-facing
        log.warning("pdf parse failed: %s", type(exc).__name__)
        raise UnsupportedFile("We couldn't open that PDF — it may be damaged.") from None


def _extract_docx(data: bytes) -> ExtractedText:
    try:
        import docx
    except ImportError:  # pragma: no cover - dependency is declared
        raise UnsupportedFile("Word support isn't available on this server.") from None

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        log.warning("docx parse failed: %s", type(exc).__name__)
        raise UnsupportedFile(
            "We couldn't open that Word file — it may be damaged."
        ) from None

    lines = [p.text for p in document.paragraphs]
    # Plenty of resumes lay their content out in tables; skipping those would
    # drop entire work histories.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("  ".join(cells))
    return ExtractedText(text="\n".join(lines))


def _tidy(text: str) -> str:
    """Normalise whitespace without destroying the line structure.

    Line breaks carry real signal in a resume — they separate a job title from
    its employer — so they survive; only runs of blank lines collapse.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Non-breaking spaces, page breaks and Unicode line separators all show up
    # in PDF text layers and confuse downstream line splitting.
    text = (text.replace("\xa0", " ")
                .replace("\x0c", "\n")
                .replace("\u2028", "\n")
                .replace("\u2029", "\n"))
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()
