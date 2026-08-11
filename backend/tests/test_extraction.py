"""Extraction: document parsing, model-output handling, and the upload API."""

import io
import json

import pytest
from fastapi.testclient import TestClient

from app.extraction.service import (ExtractionFailed, coerce, extract_resume,
                                    parse_resume_json, score_confidence)
from app.extraction.text import (ExtractedText, UnsupportedFile, extract_text,
                                 sniff_kind)
from app.llm import LlmClient, LlmError, Provider
from app.main import app
from app.ratelimit import limiter

GOOD_PAYLOAD = {
    "contact": {
        "full_name": "Sam Doe",
        "headline": "Backend Engineer",
        "email": "sam@example.com",
        "phone": "555-0100",
        "location": "Austin, TX",
        "links": [{"label": "GitHub", "url": "github.com/sam"}],
    },
    "summary": "Engineer with eight years of experience.",
    "experience": [{
        "title": "Senior Engineer", "company": "Acme", "location": "Remote",
        "start_date": "Mar 2022", "end_date": "Present",
        "bullets": ["Cut latency by 60%", "Led the migration"],
    }],
    "education": [{"degree": "B.S. CS", "institution": "UT Austin",
                   "start_date": "2015", "end_date": "2019"}],
    "skills": [{"category": "Languages", "skills": ["Python", "Go"]}],
    "projects": [],
    "publications": [],
}

RESUME_TEXT = """Sam Doe
Backend Engineer
sam@example.com | 555-0100 | Austin, TX

SUMMARY
Engineer with eight years of experience building services.

EXPERIENCE
Senior Engineer, Acme (Mar 2022 - Present)
- Cut latency by 60%
- Led the migration

EDUCATION
B.S. Computer Science, UT Austin, 2015-2019
"""


class StubClient(LlmClient):
    """Returns canned model output; records what it was asked."""

    def __init__(self, *replies):
        super().__init__([Provider(name="stub", model="stub", api_key="k",
                                   base_url="https://example.test")])
        self.replies = list(replies)
        self.prompts = []

    def complete_json(self, system, user, json_schema=None):
        self.prompts.append((system, user))
        reply = self.replies.pop(0)
        if isinstance(reply, Exception):
            raise reply
        return reply, self.providers[0]


def make_pdf(text: str) -> bytes:
    """A text-layer PDF, one source line per rendered line."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    y = 60
    for line in text.splitlines():
        page.insert_text((50, y), line, fontsize=10)
        y += 14
    return doc.tobytes()


def make_docx(paragraphs) -> bytes:
    import docx

    document = docx.Document()
    for line in paragraphs:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------- file parsing

def test_sniffs_by_content_not_extension():
    assert sniff_kind(b"%PDF-1.7 ...", "resume.docx") == "pdf"
    assert sniff_kind(b"PK\x03\x04...", "resume.pdf") == "docx"


def test_rejects_unknown_types():
    with pytest.raises(UnsupportedFile):
        extract_text(b"just some text", "resume.txt")


def test_rejects_legacy_doc_with_guidance():
    with pytest.raises(UnsupportedFile, match=r"\.docx"):
        sniff_kind(b"\xd0\xcf\x11\xe0", "resume.doc")


def test_rejects_oversized_upload():
    with pytest.raises(UnsupportedFile, match="too large"):
        extract_text(b"%PDF-" + b"x" * (6 * 1024 * 1024), "resume.pdf")


def test_reads_a_pdf():
    result = extract_text(make_pdf(RESUME_TEXT), "resume.pdf")
    assert "Sam Doe" in result.text
    assert "Acme" in result.text


def test_reads_a_docx_including_tables():
    import docx

    document = docx.Document()
    document.add_paragraph("Sam Doe")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Senior Engineer"
    table.rows[0].cells[1].text = "Acme, Mar 2022 - Present"
    document.add_paragraph("x" * 200)  # clear the minimum-text floor
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_text(buffer.getvalue(), "resume.docx")
    assert "Senior Engineer" in result.text
    assert "Acme" in result.text  # table content must not be dropped


def test_image_only_pdf_is_reported_not_silently_empty():
    import pymupdf

    doc = pymupdf.open()
    doc.new_page()  # a page with no text layer, like a scan
    with pytest.raises(UnsupportedFile, match="scan"):
        extract_text(doc.tobytes(), "scan.pdf")


# ------------------------------------------------------------- model handling

def test_parses_fenced_json():
    payload = parse_resume_json('```json\n{"a": 1}\n```')
    assert payload == {"a": 1}


def test_strips_reasoning_scratchpad():
    raw = '<think>Let me look at the header…</think>\n{"a": 1}'
    assert parse_resume_json(raw) == {"a": 1}


def test_ignores_chatter_around_the_object():
    raw = 'Sure! Here is the JSON:\n{"a": 1}\nLet me know if you need changes.'
    assert parse_resume_json(raw) == {"a": 1}


def test_coerces_invented_keys_away():
    payload = json.loads(json.dumps(GOOD_PAYLOAD))
    payload["confidence_score"] = 0.9           # a key the schema forbids
    payload["contact"]["twitter"] = "@sam"
    resume = coerce(payload)
    assert resume.contact.full_name == "Sam Doe"


def test_coerces_flat_skill_list():
    payload = json.loads(json.dumps(GOOD_PAYLOAD))
    payload["skills"] = ["Python", "Go", "SQL"]
    assert coerce(payload).skills[0].skills == ["Python", "Go", "SQL"]


def test_coerces_string_bullets_and_bare_links():
    payload = json.loads(json.dumps(GOOD_PAYLOAD))
    payload["experience"][0]["bullets"] = "Did one thing"
    payload["contact"]["links"] = ["https://github.com/sam"]
    resume = coerce(payload)
    assert resume.experience[0].bullets == ["Did one thing"]
    assert resume.contact.links[0].url == "https://github.com/sam"


def test_missing_name_gets_a_placeholder_rather_than_failing():
    payload = json.loads(json.dumps(GOOD_PAYLOAD))
    payload["contact"].pop("full_name")
    assert coerce(payload).contact.full_name == "Unknown"


def test_confidence_is_high_for_a_full_extraction():
    resume = coerce(GOOD_PAYLOAD)
    assert score_confidence(resume, ExtractedText(text=RESUME_TEXT)) >= 0.8


def test_confidence_is_low_for_a_near_empty_extraction():
    resume = coerce({"contact": {"full_name": "Sam Doe"}})
    score = score_confidence(resume, ExtractedText(text="x" * 4000))
    assert score < 0.45


# -------------------------------------------------------------- full pipeline

def test_extract_resume_end_to_end():
    client = StubClient(json.dumps(GOOD_PAYLOAD))
    result = extract_resume(make_pdf(RESUME_TEXT), "resume.pdf", client=client)
    assert result.resume.contact.full_name == "Sam Doe"
    assert result.resume.experience[0].company == "Acme"
    assert result.provider == "stub"
    # The resume text must actually reach the model.
    assert "Sam Doe" in client.prompts[0][1]


def test_invalid_output_triggers_one_repair_round():
    broken = json.dumps({"contact": {"full_name": "Sam", "email": "not-an-email"}})
    client = StubClient(broken, json.dumps(GOOD_PAYLOAD))
    result = extract_resume(make_pdf(RESUME_TEXT), "resume.pdf", client=client)
    assert result.resume.contact.email == "sam@example.com"
    assert len(client.prompts) == 2
    # The repair prompt names the offending field but never quotes its value.
    assert "contact.email" in client.prompts[1][1]
    assert "not-an-email" not in client.prompts[1][1]


def test_two_bad_replies_give_up_with_a_plain_message():
    client = StubClient("not json at all", "still not json")
    with pytest.raises(ExtractionFailed, match="fill the form in yourself"):
        extract_resume(make_pdf(RESUME_TEXT), "resume.pdf", client=client)


def test_provider_outage_is_reported_in_plain_language():
    client = StubClient(LlmError("all providers failed"))
    with pytest.raises(ExtractionFailed, match="unavailable"):
        extract_resume(make_pdf(RESUME_TEXT), "resume.pdf", client=client)


# ---------------------------------------------------------------------- API

@pytest.fixture
def client(monkeypatch):
    limiter._events.clear()
    stub = StubClient(json.dumps(GOOD_PAYLOAD))
    monkeypatch.setattr("app.main.get_client", lambda: stub)
    monkeypatch.setattr("app.extraction.service.get_client", lambda: stub)
    with TestClient(app) as c:
        yield c


def _await_job(client, job_id, timeout=30):
    import time

    deadline = time.time() + timeout
    while time.time() < deadline:
        body = client.get(f"/api/jobs/{job_id}").json()
        if body["status"] in ("done", "error"):
            return body
        time.sleep(0.1)
    raise AssertionError("job did not finish")


def test_upload_returns_the_extracted_resume(client):
    files = {"file": ("resume.pdf", make_pdf(RESUME_TEXT), "application/pdf")}
    started = client.post("/api/extract", files=files)
    assert started.status_code == 202

    body = _await_job(client, started.json()["job_id"])
    assert body["status"] == "done", body
    assert body["kind"] == "extract"
    assert body["resume"]["contact"]["full_name"] == "Sam Doe"
    assert body["low_confidence"] is False


def test_upload_of_a_scan_fails_with_guidance(client):
    import pymupdf

    doc = pymupdf.open()
    doc.new_page()
    files = {"file": ("scan.pdf", doc.tobytes(), "application/pdf")}
    started = client.post("/api/extract", files=files)
    body = _await_job(client, started.json()["job_id"])
    assert body["status"] == "error"
    assert "scan" in body["message"]
    assert body["resume"] is None


def test_upload_rejects_empty_file(client):
    files = {"file": ("resume.pdf", b"", "application/pdf")}
    assert client.post("/api/extract", files=files).status_code == 400


def test_extraction_is_rate_limited(client, monkeypatch):
    from app import main

    monkeypatch.setattr(main.EXTRACT_LIMIT, "max_events", 2, raising=False)
    files = lambda: {"file": ("resume.pdf", make_pdf(RESUME_TEXT), "application/pdf")}
    codes = [client.post("/api/extract", files=files()).status_code for _ in range(4)]
    assert 429 in codes


def test_extraction_disabled_when_no_providers(monkeypatch):
    limiter._events.clear()
    monkeypatch.setattr("app.main.get_client", lambda: LlmClient([]))
    with TestClient(app) as c:
        files = {"file": ("resume.pdf", make_pdf(RESUME_TEXT), "application/pdf")}
        response = c.post("/api/extract", files=files)
    assert response.status_code == 503
    assert "fill the form in yourself" in response.json()["error"]


def test_health_reports_providers_without_leaking_keys(client):
    body = client.get("/api/health").json()
    assert body["extraction_enabled"] is True
    assert "nvapi-" not in json.dumps(body["llm_providers"])
